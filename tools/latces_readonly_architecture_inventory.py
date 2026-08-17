#!/usr/bin/env python3
"""READ-only LAT-CES tree + SCI 1-145 evidence inventory.

Run from repository root:
    python tools/latces_readonly_architecture_inventory.py

The script reads the real checkout and SCI DOCX. It does not modify source files.
It writes only two audit outputs under audit/.
"""
from __future__ import annotations
import json, re, subprocess
from collections import defaultdict
from pathlib import Path

SCI = re.compile(r"\bSCI[\s_-]*(\d{1,3})\b", re.I)
EXTS = {".py", ".pyi", ".md", ".txt", ".rst", ".json", ".yaml", ".yml", ".toml"}

def git(*args):
    try:
        return subprocess.check_output(["git", *args], text=True, stderr=subprocess.DEVNULL).strip()
    except Exception:
        return ""

def docx_text(path: Path):
    try:
        from docx import Document
    except ImportError:
        return [], [], "python-docx is not installed"
    if not path.exists():
        return [], [], f"missing: {path}"
    d = Document(path)
    paragraphs = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    tables = []
    for ti, t in enumerate(d.tables):
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        tables.append((ti, rows))
    return paragraphs, tables, None

def scan(repo: Path):
    root = repo / "lat_ces"
    files, sci_sources = [], defaultdict(set)
    for p in root.rglob("*"):
        if not p.is_file() or "__pycache__" in p.parts:
            continue
        rel = p.relative_to(repo).as_posix()
        files.append(rel)
        if p.suffix.lower() not in EXTS:
            continue
        try:
            text = p.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            continue
        for m in SCI.finditer(text):
            n = int(m.group(1))
            if 1 <= n <= 145:
                sci_sources[n].add(rel)
    groups = defaultdict(list)
    for f in files:
        rel = f[len("lat_ces/"):] if f.startswith("lat_ces/") else f
        groups[rel.split("/", 1)[0]].append(f)
    return sorted(files), {k: sorted(v) for k, v in groups.items()}, {k: sorted(v) for k, v in sci_sources.items()}

def main():
    repo = Path(__file__).resolve().parents[1]
    out = repo / "audit"
    out.mkdir(exist_ok=True)
    files, groups, source_hits = scan(repo)
    paragraphs, tables, docx_error = docx_text(repo / "SCI 1-145 LAT SES.docx")
    doc_hits = defaultdict(list)
    for i, text in enumerate(paragraphs, 1):
        for m in SCI.finditer(text):
            n = int(m.group(1))
            if 1 <= n <= 145:
                doc_hits[n].append({"location": f"paragraph:{i}", "text": text})
    for ti, rows in tables:
        for ri, row in enumerate(rows, 1):
            text = " | ".join(row)
            for m in SCI.finditer(text):
                n = int(m.group(1))
                if 1 <= n <= 145:
                    doc_hits[n].append({"location": f"table:{ti}:row:{ri}", "text": text})
    mapping = {}
    for n in range(1, 146):
        d, s = doc_hits[n], source_hits.get(n, [])
        mapping[str(n)] = {"status": "MAPPED" if d and s else "DOCUMENT_ONLY" if d else "SOURCE_ONLY" if s else "UNMAPPED", "document": d, "repository_sources": s}
    data = {"branch": git("branch", "--show-current"), "commit": git("rev-parse", "HEAD"), "lat_ces_files_excluding_pycache": len(files), "groups": groups, "docx_error": docx_error, "sci_1_145": mapping}
    (out / "LAT_CES_READONLY_MAIN_INVENTORY.json").write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    md = ["# LAT-CES READ-ONLY MAIN inventory", "", f"- branch: `{data['branch']}`", f"- commit: `{data['commit']}`", f"- lat_ces files: **{len(files)}**", f"- SCI DOCX: **{'OK' if not docx_error else docx_error}**", "", "## Tree"]
    for g, fs in sorted(groups.items()):
        md += [f"### lat_ces/{g} ({len(fs)})"] + [f"- `{f}`" for f in fs]
    md += ["", "## SCI 1-145 evidence"]
    for n in range(1, 146):
        x = mapping[str(n)]
        md.append(f"### SCI {n} — **{x['status']}**")
        md += [f"- DOC {a['location']}: {a['text']}" for a in x['document']]
        md += [f"- SRC `{s}`" for s in x['repository_sources']]
        if not x['document'] and not x['repository_sources']:
            md.append("- No evidence found by this scan.")
    (out / "LAT_CES_READONLY_MAIN_INVENTORY.md").write_text("\n".join(md) + "\n", encoding="utf-8")
    print(f"branch={data['branch']} commit={data['commit']}")
    print(f"lat_ces_files={len(files)}")
    print("outputs: audit/LAT_CES_READONLY_MAIN_INVENTORY.md and .json")

if __name__ == "__main__":
    main()
